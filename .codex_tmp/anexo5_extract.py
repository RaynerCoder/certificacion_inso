from pathlib import Path
import json
import sys

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def main():
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    document = Document(source)
    blocks = []

    for index, block in enumerate(iter_block_items(document), start=1):
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if text:
                blocks.append({
                    "index": index,
                    "type": "paragraph",
                    "style": block.style.name if block.style else None,
                    "text": text,
                })
        else:
            rows = []
            for row in block.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            blocks.append({
                "index": index,
                "type": "table",
                "rows": rows,
            })

    payload = {
        "source": str(source),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
        "blocks": blocks,
    }
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({key: payload[key] for key in ("source", "paragraphs", "tables", "inline_shapes", "sections")}, ensure_ascii=False))


if __name__ == "__main__":
    main()
