from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE = Path(
    r"C:\Users\Rayner\Downloads\Material Sistema Certificador"
    r"\Guia flujo empresa y tramitadores - INSO.docx"
)
OUTPUT = Path(
    r"C:\laragon\www\sistema_certificador_inso\.codex_tmp\producto_guia"
    r"\Guia flujo empresa tramitadores y productos - INSO.docx"
)

NAVY = "12304A"
TEAL = "00897B"
TEAL_LIGHT = "E6F5F2"
BLUE_LIGHT = "EAF2F8"
AMBER = "C47A00"
AMBER_LIGHT = "FFF4D6"
RED = "B42318"
RED_LIGHT = "FDECEC"
GREEN = "147D64"
GREEN_LIGHT = "E8F6F1"
GRAY = "556372"
GRAY_LIGHT = "F2F5F7"
WHITE = "FFFFFF"
BORDER = "CAD5DF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=110, start=130, bottom=110, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, *, bold=False, color=NAVY, size=8.6, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def style_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        if r_idx == 0 and header:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            if widths and c_idx < len(widths):
                cell.width = widths[c_idx]
            if r_idx == 0 and header:
                shade(cell, NAVY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            elif r_idx % 2 == 0:
                shade(cell, GRAY_LIGHT)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": BORDER},
                left={"val": "single", "sz": 4, "color": BORDER},
                bottom={"val": "single", "sz": 4, "color": BORDER},
                right={"val": "single", "sz": 4, "color": BORDER},
            )


def add_body(doc, text, *, bold_prefix=None, space_after=4, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.color.rgb = RGBColor.from_string(GRAY)
    else:
        r = p.add_run(text)
        r.font.color.rgb = RGBColor.from_string(GRAY)
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9.3)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(9.1)
        r.font.color.rgb = RGBColor.from_string(GRAY)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.72)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(2.8)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(9.1)
        r.font.color.rgb = RGBColor.from_string(GRAY)


def add_callout(doc, title, body, *, fill=TEAL_LIGHT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    shade(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 5, "color": accent},
        left={"val": "single", "sz": 18, "color": accent},
        bottom={"val": "single", "sz": 5, "color": accent},
        right={"val": "single", "sz": 5, "color": accent},
    )
    cell_margins(cell, top=150, start=190, bottom=150, end=190)
    cell.text = ""
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(3)
    r = p_title.add_run(title)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(accent)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.1
    r = p_body.add_run(body)
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_flow(doc, steps):
    table = doc.add_table(rows=len(steps), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (title, detail) in enumerate(steps, start=1):
        num_cell, text_cell = table.rows[index - 1].cells
        num_cell.width = Inches(0.58)
        text_cell.width = Inches(6.12)
        shade(num_cell, TEAL if index < len(steps) else NAVY)
        shade(text_cell, WHITE if index % 2 else GRAY_LIGHT)
        set_cell_text(num_cell, index, bold=True, color=WHITE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        text_cell.text = ""
        p = text_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Aptos"
        r.font.size = Pt(9.2)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p2 = text_cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.05
        r2 = p2.add_run(detail)
        r2.font.name = "Aptos"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor.from_string(GRAY)
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(text_cell, top=90, start=140, bottom=90, end=140)
        for cell in (num_cell, text_cell):
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": BORDER},
                left={"val": "single", "sz": 4, "color": BORDER},
                bottom={"val": "single", "sz": 4, "color": BORDER},
                right={"val": "single", "sz": 4, "color": BORDER},
            )
        prevent_row_split(table.rows[index - 1])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE, size=8.3)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, color=NAVY if i == 0 else GRAY, size=8.2)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document(SOURCE)

# Portada y alcance: se amplían sin alterar las conclusiones originales.
doc.paragraphs[1].text = "Empresa, tramitadores y productos"
doc.paragraphs[2].text = "Flujos actuales, efectos de la baja y relación de productos con trámites"
doc.paragraphs[6].text = (
    "Explica el comportamiento comprobado en controladores, servicios, modelos, rutas, vistas y migraciones."
)
doc.paragraphs[8].text = (
    "Incluye procedimientos operativos para retirar a un tramitador y para registrar productos dentro o fuera de un trámite."
)
for idx in (1, 2, 6, 8):
    for run in doc.paragraphs[idx].runs:
        run.font.name = "Aptos"

# Actualiza el encabezado de las páginas interiores.
for section in doc.sections:
    for paragraph in section.header.paragraphs:
        if "GUIA EMPRESA Y TRAMITADORES" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(
                    "GUIA EMPRESA Y TRAMITADORES",
                    "GUIA EMPRESA, TRAMITADORES Y PRODUCTOS",
                )

doc.core_properties.title = "Guía de empresa, tramitadores y productos - INSO"
doc.core_properties.subject = "Flujo funcional comprobado en el Sistema Certificador INSO"

doc.add_page_break()
add_heading(doc, "10. Respuesta directa sobre el registro de productos", 1)
add_callout(
    doc,
    "Diferencia principal",
    "Registrar desde el módulo Productos crea el catálogo del producto, pero no lo asocia a un trámite. "
    "Registrar desde el detalle de un trámite en revisión crea el producto y sus registros, y después enlaza "
    "esos registros al certificado mediante la tabla certificados_registros.",
)
add_body(
    doc,
    "La relación con el trámite no se realiza directamente desde productos. El certificado se conecta con uno o "
    "varios registros; cada registro identifica el producto y la presentación utilizada.",
)

add_heading(doc, "10.1. Estructura real de la relación", 2)
add_flow(
    doc,
    [
        ("Persona beneficiaria o importadora", "Se guarda en productos.id_importador_persona."),
        ("Producto", "Contiene código, nombres, país, fabricante, tipo, clasificación y estado."),
        ("Presentaciones y registros", "Las presentaciones describen formato y etiqueta; los registros guardan autorización, vigencia y cantidad."),
        ("Tabla certificados_registros", "Relaciona cada registro creado con el certificado que originó la operación."),
        ("Detalle del trámite", "Agrupa los registros por producto y muestra los productos asociados."),
    ],
)
add_callout(
    doc,
    "Los identificadores no se reinician",
    "Todas estas tablas conservan su id autoincremental. El flujo inserta filas nuevas y usa las llaves foráneas; "
    "no trunca tablas ni reinicia AUTO_INCREMENT.",
    fill=BLUE_LIGHT,
    accent=NAVY,
)

add_heading(doc, "11. Flujo A: registrar un producto por separado", 1)
add_body(
    doc,
    "Este flujo comienza en Menú > Productos > Nuevo producto. Según los permisos predeterminados, puede acceder "
    "el personal que tenga productos.ver; los roles externos Solicitante, Representante legal y Tramitador no lo "
    "reciben por defecto.",
)
add_numbers(
    doc,
    [
        "El usuario abre Productos y selecciona Nuevo producto.",
        "Completa importador, código, nombre comercial, país, fabricante, tipo, clasificación opcional y estado.",
        "Agrega ingredientes si corresponde. Cada ingrediente se relaciona con el producto y puede incluir porcentaje.",
        "Agrega presentaciones si corresponde. Toda presentación nueva exige una etiqueta en PDF de hasta 5 MB.",
        "Puede agregar registros de autorización, pero en este flujo no son obligatorios.",
        "Al guardar, el sistema ejecuta toda la operación dentro de una transacción de base de datos.",
        "Si todo es válido, vuelve al listado general de Productos.",
    ],
)
add_matrix(
    doc,
    ["Elemento", "¿Se guarda?", "Condición"],
    [
        ("productos", "Sí", "Siempre se crea una fila nueva."),
        ("ingredientes_productos", "Opcional", "Solo por cada ingrediente agregado."),
        ("presentaciones", "Opcional", "Solo si se agregan presentaciones nuevas."),
        ("registros", "Opcional", "Solo si se agregan autorizaciones."),
        ("certificados_registros", "No", "No existe form_id_certificado en el origen."),
        ("Archivo de etiqueta", "Opcional", "Se guarda en storage/app/public/productos/etiquetas."),
    ],
    [Inches(1.55), Inches(1.1), Inches(4.05)],
)
add_callout(
    doc,
    "Resultado del flujo separado",
    "El producto queda disponible en el catálogo general, pero no aparecerá en “Productos asociados” de ningún "
    "trámite. Tener un registro de autorización tampoco crea automáticamente la relación con un certificado.",
    fill=AMBER_LIGHT,
    accent=AMBER,
)

add_heading(doc, "11.1. Qué puede hacerse después", 2)
add_bullets(
    doc,
    [
        "El listado muestra todos los productos al usuario que tenga productos.ver; actualmente no filtra por empresa o persona importadora.",
        "El sistema puede reutilizar una presentación existente al crear otro producto si coincide el importador y el código; si no hay código, compara el nombre comercial.",
        "Las rutas Ver, Editar y Eliminar existen y los botones se muestran, pero los métodos correspondientes del controlador Producto están vacíos. En el código actual solo están completos el listado y el registro.",
    ],
)

doc.add_page_break()
add_heading(doc, "12. Flujo B: registrar desde Revisar trámite", 1)
add_body(
    doc,
    "El botón Registrar producto aparece en el detalle técnico únicamente cuando el tipo de certificado requiere "
    "evidencia PRODUCTO, el usuario no es el solicitante, puede asignar o revisar requisitos y el trámite tiene "
    "una persona beneficiaria.",
)
add_flow(
    doc,
    [
        ("Abrir el trámite asignado", "El funcionario entra al detalle y revisa los requisitos configurados."),
        ("Abrir Registrar producto", "La URL lleva el id del certificado, el id del beneficiario y la pantalla de retorno."),
        ("Completar el producto", "Se usan las mismas cuatro etapas del módulo Productos."),
        ("Agregar al menos un registro", "Cuando existe form_id_certificado, el backend exige por lo menos un registro de autorización."),
        ("Validar el beneficiario", "El producto solo puede guardarse para la misma Persona beneficiaria del trámite."),
        ("Crear y relacionar", "Se crea el producto y se agregan los registros nuevos a certificados_registros sin quitar relaciones previas."),
        ("Volver al detalle", "Los registros se agrupan por producto dentro de Productos asociados."),
        ("Revisar el requisito", "El técnico todavía debe marcar manualmente si el requisito PRODUCTO cumple o no cumple."),
    ],
)

add_heading(doc, "12.1. Datos que se transportan desde el trámite", 2)
add_matrix(
    doc,
    ["Dato", "Uso real"],
    [
        ("form_id_certificado", "Identifica el certificado al que se enlazarán los registros nuevos."),
        ("form_id_importador_persona", "Preselecciona al beneficiario; el servidor exige que coincida con el trámite."),
        ("bandeja", "Conserva el contexto de la bandeja de trabajo."),
        ("return_to", "Devuelve al detalle del trámite; solo se aceptan direcciones internas."),
    ],
    [Inches(2.05), Inches(4.65)],
)
add_callout(
    doc,
    "La revisión no es automática",
    "Crear y asociar el producto no cambia requisitos_certificados.cumple ni el estado del requisito. El técnico "
    "debe comprobar la información mostrada y guardar su decisión de revisión.",
    fill=BLUE_LIGHT,
    accent=NAVY,
)

add_heading(doc, "13. Qué ocurre al iniciar un trámite que exige PRODUCTO", 1)
add_numbers(
    doc,
    [
        "El solicitante inicia el trámite y el sistema crea el requisito con estado PENDIENTE_REVISION.",
        "Para evidencia PRODUCTO no se exige un archivo PDF ni una imagen durante el registro inicial.",
        "El flujo inicial actual no enlaza un producto existente ni crea una fila en certificados_registros.",
        "Cuando el trámite llega al detalle técnico, el funcionario autorizado puede registrar el producto desde Productos asociados.",
        "Después de la asociación, el técnico revisa el requisito y decide SI o NO.",
        "El trámite solo puede finalizar cuando todos los requisitos, incluido PRODUCTO, hayan sido marcados SI.",
    ],
)
add_body(
    doc,
    "Existe código de interfaz para mostrar el formulario de producto en modo embebido y enviar un resumen visual "
    "a una pantalla contenedora. Sin embargo, en la implementación revisada no hay una vista de Nuevo trámite que "
    "lo invoque ni lógica que use ese resumen para enlazar registros al certificado. Por ello, no debe considerarse "
    "un flujo operativo completo.",
)

add_heading(doc, "14. Comparación de los dos caminos", 1)
add_matrix(
    doc,
    ["Aspecto", "Productos, por separado", "Desde Revisar trámite"],
    [
        ("Origen", "Menú Productos.", "Detalle técnico del certificado."),
        ("Permiso", "productos.ver.", "Permisos de revisión y, además, productos.ver por la ruta."),
        ("Beneficiario", "El usuario elige una Persona activa.", "Debe coincidir con el beneficiario del trámite."),
        ("Registro", "Opcional.", "Al menos uno es obligatorio."),
        ("Vínculo al trámite", "No se crea.", "Sí, mediante certificados_registros."),
        ("Pantalla final", "Listado general.", "Regresa al detalle del trámite."),
        ("Cumplimiento", "No afecta requisitos.", "No se marca automáticamente; decide el técnico."),
    ],
    [Inches(1.35), Inches(2.65), Inches(2.7)],
)

doc.add_page_break()
add_heading(doc, "15. Si el producto ya fue registrado antes", 1)
add_callout(
    doc,
    "Comportamiento actual",
    "El formulario no permite seleccionar un producto existente y asociar directamente uno de sus registros al "
    "trámite. Cada guardado ejecuta Producto::create, por lo que genera una nueva fila en productos.",
    fill=AMBER_LIGHT,
    accent=AMBER,
)
add_bullets(
    doc,
    [
        "No existe una regla única de base de datos ni una validación que impida repetir código o nombre comercial.",
        "El sistema sí puede reutilizar una presentación existente del mismo importador y del mismo código o nombre, pero el registro nuevo se guarda con el producto recién creado.",
        "Una presentación reutilizada conserva su id_producto original. Por ello, el registro nuevo puede apuntar a un producto nuevo y, al mismo tiempo, a una presentación perteneciente al producto anterior.",
        "La tabla certificados_registros solo recibe los registros creados en esa operación; un registro creado antes por separado no se selecciona ni se adjunta desde esta pantalla.",
    ],
)
add_callout(
    doc,
    "Uso operativo mientras no se cambie el código",
    "Si el producto debe formar parte de un trámite, conviene registrarlo desde ese trámite y verificar antes que no "
    "exista otro con el mismo código. Registrar primero por separado y volver a registrarlo desde revisión puede "
    "duplicar la fila de productos.",
    fill=GREEN_LIGHT,
    accent=GREEN,
)

add_heading(doc, "16. Flujo cuando el requisito PRODUCTO es observado", 1)
add_body(
    doc,
    "El técnico puede marcar el requisito como NO, escribir una observación y notificar la corrección al "
    "beneficiario o a un tramitador autorizado. En la pantalla de corrección se muestra un enlace Registrar producto.",
)
add_matrix(
    doc,
    ["Punto revisado", "Comportamiento vigente"],
    [
        ("Acceso al enlace", "La ruta exige productos.ver."),
        ("Roles externos predeterminados", "Solicitante, Representante legal y Tramitador no tienen productos.ver."),
        ("Efecto posible", "Con permisos predeterminados, el enlace puede responder 403."),
        ("Validación al reenviar", "Exige corrección nueva para PDF y TEXTO; no exige una asociación nueva para PRODUCTO."),
        ("Estado al reenviar", "El requisito vuelve a PENDIENTE_REVISION y el técnico debe revisarlo otra vez."),
    ],
    [Inches(2.1), Inches(4.6)],
)
add_callout(
    doc,
    "Inconsistencia funcional encontrada",
    "La interfaz ofrece Registrar producto al responsable de la corrección, pero la autorización predeterminada "
    "puede impedirle entrar. Además, el backend permite reenviar la corrección PRODUCTO sin comprobar que se haya "
    "creado o asociado un registro nuevo.",
    fill=RED_LIGHT,
    accent=RED,
)

add_heading(doc, "17. Controles y riesgos encontrados", 1)
add_matrix(
    doc,
    ["Control o riesgo", "Situación actual", "Impacto"],
    [
        ("Tipo de trámite", "Se valida que el certificado realmente requiera PRODUCTO.", "Evita asociaciones fuera de configuración."),
        ("Beneficiario", "Se valida que producto e trámite usen la misma Persona.", "Evita mezclar importadores."),
        ("Registro mínimo", "Es obligatorio al registrar desde trámite.", "Permite crear la relación técnica."),
        ("Acceso al certificado", "ProductoController no comprueba que el usuario esté asignado o autorizado para ese certificado concreto.", "Un usuario con productos.ver podría intentar usar otro id de certificado."),
        ("Duplicidad de producto", "No hay unicidad por código, nombre o importador.", "Puede duplicarse el catálogo."),
        ("Cumplimiento", "La asociación no actualiza automáticamente el requisito.", "Depende de la revisión manual."),
        ("Desvinculación", "No hay una acción de interfaz para quitar una relación producto-trámite.", "Una asociación incorrecta requiere intervención técnica."),
    ],
    [Inches(1.55), Inches(3.15), Inches(2.0)],
)

add_heading(doc, "18. Tablas y campos que se llenan", 1)
add_matrix(
    doc,
    ["Tabla", "Campos principales", "Cuándo"],
    [
        ("productos", "id autoincremental; importador, país, fabricante, tipo, código, nombres, clasificación, estado y auditoría.", "Siempre."),
        ("ingredientes_productos", "id; id_producto, id_ingrediente, porcentaje, estado y auditoría.", "Si se agregan ingredientes."),
        ("presentaciones", "id; id_producto, etiqueta, cantidad, unidad, descripción, estado y auditoría.", "Si la presentación es nueva."),
        ("registros", "id; id_producto, autorización, vigencia, cantidad, unidad, id_presentacion, estado y auditoría.", "Si se agregan registros; obligatorio desde trámite."),
        ("certificados_registros", "id; id_certificado, id_registro, auditoría y fechas.", "Solo cuando el formulario lleva form_id_certificado."),
        ("Catálogos relacionados", "fabricantes, tipos_productos, ingredientes, clasificaciones_productos o catalogos_medidas.", "Solo si el usuario crea una opción temporal nueva."),
    ],
    [Inches(1.45), Inches(3.75), Inches(1.5)],
)

add_heading(doc, "19. Archivos adicionales revisados", 1)
add_bullets(
    doc,
    [
        "app/Http/Controllers/ProductoController.php",
        "app/Http/Controllers/SeguimientoController.php",
        "app/Http/Middleware/VerificarPermiso.php",
        "app/Livewire/Datatables/ProductoTable.php",
        "app/Models/Producto.php, Registro.php, Presentacion.php y Certificado.php",
        "resources/views/productos/create.blade.php",
        "resources/views/productos/create/*.blade.php",
        "resources/views/seguimientos_certificados/producto/*.blade.php",
        "resources/views/certificados/componentes/detalle_tramite.blade.php",
        "routes/admin.php y bootstrap/app.php",
        "migraciones de productos, ingredientes_productos, presentaciones, registros y certificados_registros",
        "seeders PermisoSeeder.php y PermisoRoleSeeder.php",
    ],
)

add_callout(
    doc,
    "Conclusión",
    "Para el funcionamiento actual, el camino que realmente enlaza un producto a un trámite es Registrar producto "
    "desde el detalle del trámite y guardar por lo menos un registro. El alta separada sirve como catálogo, pero "
    "no reemplaza esa relación. Los vacíos de permisos, reutilización y validación señalados deben considerarse "
    "antes de usar el flujo en producción.",
    fill=TEAL_LIGHT,
    accent=TEAL,
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
