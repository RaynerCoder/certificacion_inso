from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "guia_flujo_empresa_tramitadores_inso.docx"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "17365D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
TEAL = "0F766E"
GREEN = "15803D"
AMBER = "B45309"
RED = "B42318"
SLATE = "475569"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F5F3"
LIGHT_AMBER = "FFF4E5"
LIGHT_RED = "FDECEC"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
BORDER = "CBD5E1"
TEXT = "1F2937"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_table_fixed_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_run_style(run, size=11, color=TEXT, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, bold=False, italic=False, color=TEXT, size=11):
    run = paragraph.add_run(text)
    set_run_style(run, size=size, color=color, bold=bold, italic=italic)
    return run


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_style(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_running_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    add_text(p, "SISTEMA CERTIFICADOR - INSO", bold=True, color=NAVY, size=8.5)
    add_text(p, "  |  GUIA EMPRESA Y TRAMITADORES", color=MUTED, size=8.5)

    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0]
    add_page_number(p_footer)
    set_paragraph_spacing(p_footer, after=0, line=1.0)


def add_numbering_definition(document, fmt="decimal", text="%1.", color=TEXT):
    numbering = document.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = (max(existing_abs) + 1) if existing_abs else 100

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = (max(existing_num) + 1) if existing_num else 100
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_list_item(document, text, num_id, bold_prefix=None, color=TEXT):
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    set_paragraph_spacing(p, after=4)
    p.paragraph_format.keep_together = True

    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=color)
        add_text(p, text[len(bold_prefix):], color=color)
    else:
        add_text(p, text, color=color)
    return p


def add_body(document, text, bold_prefix=None, color=TEXT, after=6):
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=after)
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=color)
        add_text(p, text[len(bold_prefix):], color=color)
    else:
        add_text(p, text, color=color)
    return p


def add_heading(document, text, level=1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(document, title, body, fill, accent):
    table = document.add_table(rows=1, cols=1)
    set_table_fixed_layout(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=accent, size="8")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_text(p, title, bold=True, color=accent, size=11)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0)
    add_text(p2, body, color=TEXT)
    after = document.add_paragraph()
    set_paragraph_spacing(after, after=2, line=1.0)
    return table


def add_content_table(document, headers, rows, widths, header_fill=LIGHT_BLUE):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table = document.add_table(rows=1, cols=len(headers))
    set_table_fixed_layout(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.1)
        add_text(p, label, bold=True, color=NAVY, size=9.5)

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFBFC")
            p = cells[index].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            add_text(p, str(value), size=9.3)
    set_table_fixed_layout(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True


def add_cover(document):
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.0)
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("GUIA OPERATIVA")
    set_run_style(run, size=10, color=TEAL, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    set_paragraph_spacing(p, before=42, after=8, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Empresa y tramitadores", bold=True, color=NAVY, size=28)

    p = document.add_paragraph()
    set_paragraph_spacing(p, after=22, line=1.1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Flujo actual, efectos de la baja y controles seguros", color=SLATE, size=15)

    p = document.add_paragraph()
    set_paragraph_spacing(p, after=30, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Sistema Certificador - INSO", bold=True, color=BLUE_DARK, size=11)
    add_text(p, "\nRevisión del código vigente: 30 de julio de 2026", color=MUTED, size=10)

    add_callout(
        document,
        "Respuesta directa",
        "Si la empresa da de baja correctamente a un tramitador desde el módulo Tramitadores, "
        "la persona y los trámites históricos no se borran. Se desactiva únicamente la relación "
        "con esa empresa; las correcciones abiertas asignadas a ese tramitador pasan a la cuenta "
        "activa de la empresa beneficiaria.",
        LIGHT_TEAL,
        TEAL,
    )

    p = document.add_paragraph()
    set_paragraph_spacing(p, before=12, after=5)
    add_text(p, "Alcance del documento", bold=True, color=NAVY, size=11)
    bullets = add_numbering_definition(document, fmt="bullet", text="•", color=TEAL)
    for item in [
        "Explica el comportamiento comprobado en controladores, servicios, modelos y rutas.",
        "Distingue la baja segura de otras acciones que parecen similares, pero no producen el mismo resultado.",
        "Incluye un procedimiento operativo antes, durante y después de retirar a un tramitador.",
        "No modifica el sistema ni propone que se borren datos históricos.",
    ]:
        add_list_item(document, item, bullets)

    p = document.add_paragraph()
    set_paragraph_spacing(p, before=18, after=0, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Documento de orientación funcional y técnica", italic=True, color=MUTED, size=9.5)


def build_document():
    document = Document()
    configure_styles(document)
    configure_page(document.sections[0])
    add_cover(document)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(content_section)
    set_page_number_start(content_section, 1)
    add_running_header(content_section)

    add_heading(document, "1. Cómo representa el sistema a la empresa y al tramitador", 1)
    add_body(
        document,
        "En el sistema, una empresa, un tramitador y una cuenta de acceso son registros relacionados, "
        "pero no son el mismo registro. Esta separación permite conservar el historial aunque una relación "
        "operativa deje de estar activa.",
    )
    add_content_table(
        document,
        ["Elemento", "Qué representa y dónde se usa"],
        [
            ("Persona de la empresa", "Es la identidad base del beneficiario. Puede estar vinculada a una cuenta de usuario."),
            ("Empresa", "Contiene la razón social y los datos empresariales; pertenece a la Persona de la empresa."),
            ("Persona natural", "Es la identidad base del tramitador, con sus datos naturales, teléfonos y rubros."),
            ("Responsable", "Une a la Persona natural con una Empresa y un rol. Para este flujo, el rol es tramitador."),
            ("Usuario", "Permite iniciar sesión, iniciar trámites o recibir correcciones. Se vincula a una Persona."),
            ("Certificado / trámite", "Guarda por separado al beneficiario y al tramitador que intervino en la solicitud."),
            ("Seguimiento", "Guarda cada movimiento y el usuario al que corresponde el siguiente paso."),
        ],
        [2500, 6860],
    )
    add_callout(
        document,
        "Idea clave",
        "Dar de baja al tramitador no significa borrar su Persona, su Usuario ni los trámites. "
        "Significa desactivar una relación Responsable entre esa Persona y una Empresa concreta.",
        LIGHT_BLUE,
        BLUE,
    )

    add_heading(document, "2. Flujo completo paso a paso", 1)
    add_heading(document, "2.1. La empresa existe como beneficiaria", 2)
    steps = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "La empresa tiene una Persona base y un registro Empresa asociado.",
        "La cuenta principal de la empresa, cuando existe y está activa, se vincula a esa Persona mediante id_usuario.",
        "En un trámite de la empresa, la Persona de la empresa queda registrada como beneficiaria.",
    ]:
        add_list_item(document, item, steps)

    add_heading(document, "2.2. La empresa registra un tramitador", 2)
    steps = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "La cuenta de la empresa entra al módulo Tramitadores; el acceso exige el permiso tramitadores.ver.",
        "El sistema crea una Persona y una ficha Natural para el tramitador.",
        "También registra sus teléfonos, sus rubros CAEB seleccionados y, si corresponde, el respaldo PDF.",
        "Finalmente crea el Responsable que relaciona a la Persona natural con la Empresa, con rol tramitador, fecha de registro y estado.",
    ]:
        add_list_item(document, item, steps)
    add_callout(
        document,
        "Limitación comprobada",
        "El registro realizado por TramitadorController no crea una cuenta User ni la enlaza automáticamente. "
        "Por tanto, una persona puede aparecer como tramitador administrativo, pero no podrá iniciar sesión, "
        "iniciar trámites ni recibir correcciones hasta tener un Usuario activo vinculado a su Persona.",
        LIGHT_AMBER,
        AMBER,
    )

    add_heading(document, "2.3. El tramitador opera con una cuenta activa", 2)
    bullets = add_numbering_definition(document, fmt="bullet", text="•", color=TEAL)
    for item in [
        "La Persona del tramitador debe estar activa.",
        "Debe existir una relación Responsable activa con la empresa y con rol tramitador.",
        "Para operar en pantalla debe existir un Usuario activo vinculado a esa Persona.",
        "El Usuario debe tener los permisos del rol necesarios para iniciar y consultar trámites.",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "2.4. Se inicia el trámite", 2)
    add_content_table(
        document,
        ["Quién lo inicia", "Qué registra el servidor"],
        [
            (
                "Cuenta de la empresa",
                "La empresa es beneficiaria. Como la cuenta externa no puede cambiar al tramitador desde el navegador, "
                "el servidor usa la Persona titular de la cuenta como tramitador del trámite.",
            ),
            (
                "Tramitador activo",
                "Puede elegir como beneficiaria a su propia Persona y a las empresas donde tiene una relación activa. "
                "El servidor fuerza como tramitador a la Persona de la cuenta autenticada.",
            ),
            (
                "Personal interno autorizado",
                "Puede seleccionar un tramitador activo perteneciente a la empresa. El servidor valida esa relación antes de guardar.",
            ),
        ],
        [2700, 6660],
    )
    add_body(
        document,
        "El trámite conserva dos datos históricos distintos: id_persona_beneficiario e id_persona_tramitador. "
        "La baja posterior del tramitador no elimina ni reemplaza esos valores.",
    )

    add_heading(document, "2.5. El trámite avanza y puede ser observado", 2)
    steps = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "El trámite se registra y el Seguimiento indica qué usuario debe atender el siguiente paso.",
        "Si el personal técnico observa requisitos, la corrección puede asignarse al beneficiario o a un tramitador activo que tenga Usuario activo.",
        "Como destinatario predeterminado se prefiere al tramitador original si todavía está habilitado; si no, se prefiere al beneficiario.",
        "La persona que responde debe seguir autorizada en el momento de la respuesta.",
        "Al finalizar, el historial del trámite sigue mostrando quién fue el beneficiario y quién actuó como tramitador.",
    ]:
        add_list_item(document, item, steps)

    document.add_page_break()
    add_heading(document, "3. Qué ocurre al dar de baja al tramitador", 1)
    add_callout(
        document,
        "Ruta correcta",
        "Empresa autenticada > Tramitadores > seleccionar el registro > Dar de baja. "
        "El sistema ejecuta la baja dentro de una transacción para que todos los cambios se apliquen juntos o ninguno se aplique.",
        LIGHT_TEAL,
        TEAL,
    )

    add_heading(document, "3.1. Proceso interno de la baja", 2)
    steps = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "Comprueba que el Responsable seleccionado pertenece a la empresa autenticada.",
        "Comprueba que todavía es un tramitador activo.",
        "Busca correcciones pendientes asignadas a su Usuario: Seguimiento activo, sin fecha de derivación y con trámite en estado OBSERVADO.",
        "Si existe una corrección pendiente, comprueba que la empresa beneficiaria tenga un Usuario activo para recibirla.",
        "Cierra el movimiento anterior colocando la fecha de derivación.",
        "Crea un nuevo Seguimiento hijo asignado al Usuario de la empresa beneficiaria, con la referencia Baja del tramitador.",
        "Crea una notificación para la empresa y desactiva las notificaciones activas del tramitador relacionadas con esa empresa.",
        "Actualiza el Responsable a INACTIVO y registra la fecha de baja.",
        "Informa cuántos trámites pendientes fueron transferidos.",
    ]:
        add_list_item(document, item, steps)

    add_heading(document, "3.2. Qué cambia y qué se conserva", 2)
    add_content_table(
        document,
        ["Dato u objeto", "Resultado después de la baja"],
        [
            ("Relación Responsable", "Cambia a INACTIVO y recibe fecha_baja."),
            ("Persona del tramitador", "Se conserva. No se borra ni se inactiva automáticamente."),
            ("Ficha Natural, teléfonos y rubros", "Se conservan."),
            ("Usuario del tramitador", "Se conserva y no se desactiva automáticamente."),
            ("Trámites históricos", "Se conservan con el tramitador original registrado."),
            ("Correcciones OBSERVADAS abiertas", "Pasan a la cuenta activa de la empresa beneficiaria."),
            ("Otros trámites en revisión o finalizados", "No se reasignan; continúan y mantienen el historial."),
            ("Notificaciones activas de esa empresa", "Las destinadas al tramitador se marcan INACTIVAS."),
            ("Otros tramitadores de la empresa", "No cambian; continúan activos."),
        ],
        [3020, 6340],
    )

    add_heading(document, "3.3. Efecto para el tramitador dado de baja", 2)
    bullets = add_numbering_definition(document, fmt="bullet", text="•", color=TEAL)
    for item in [
        "Ya no puede iniciar nuevos trámites para esa empresa.",
        "Ya no puede responder correcciones de esa empresa.",
        "Pierde la autorización de detalle que dependía de la relación activa con esa empresa.",
        "No pierde automáticamente el acceso a otras empresas donde conserve otra relación activa.",
        "No desaparece del historial de los trámites que gestionó.",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "3.4. Cuándo el sistema rechaza la baja", 2)
    add_body(
        document,
        "Si hay una corrección pendiente y la empresa beneficiaria no tiene una cuenta User activa, la baja se rechaza. "
        "Como la operación se ejecuta en una transacción, no se debe cerrar parcialmente el seguimiento ni desactivar parcialmente la relación.",
    )

    add_heading(document, "4. Escenarios frecuentes", 1)
    add_content_table(
        document,
        ["Escenario", "Qué ocurre"],
        [
            ("Es el único tramitador y no tiene correcciones pendientes", "Se da de baja la relación. La empresa queda sin tramitador activo, pero los trámites históricos permanecen."),
            ("Tiene correcciones OBSERVADAS pendientes", "Cada corrección abierta se transfiere al Usuario activo de la empresa beneficiaria."),
            ("La empresa no tiene Usuario activo", "Si hay correcciones por transferir, la baja se bloquea hasta resolver el acceso del beneficiario."),
            ("La empresa tiene varios tramitadores", "Solo se desactiva el seleccionado. Los demás continúan activos."),
            ("La misma persona tramita para varias empresas", "Solo cambia la relación con la empresa que ejecuta la baja. Las demás relaciones no se modifican."),
            ("El tramitador no tiene Usuario", "Puede existir como registro administrativo, pero no puede operar. Tampoco puede tener una corrección asignada por id_usuario."),
            ("El trámite ya finalizó", "No se modifica. El tramitador original permanece visible para auditoría."),
            ("Se registra un reemplazo", "El reemplazo debe quedar activo y con Usuario. La baja del anterior transfiere pendientes al beneficiario, no automáticamente al reemplazo."),
        ],
        [3280, 6080],
    )

    document.add_page_break()
    add_heading(document, "5. No todas las formas de “eliminar” son iguales", 1)
    add_content_table(
        document,
        ["Acción", "Comportamiento actual", "Evaluación"],
        [
            (
                "Dar de baja en Tramitadores",
                "Transfiere correcciones pendientes, conserva historial y desactiva solo la relación.",
                "Ruta segura y recomendada.",
            ),
            (
                "Eliminar desde Responsables",
                "El controlador bloquea la eliminación cuando el Responsable es tramitador y exige usar Tramitadores.",
                "Protección correcta.",
            ),
            (
                "Quitar al editar la ficha Persona de la empresa",
                "La actualización elimina lógicamente las relaciones Responsables y las vuelve a crear desde el formulario; no ejecuta la transferencia de correcciones.",
                "Evitar como método de baja.",
            ),
            (
                "Inactivar o eliminar el Usuario",
                "Desactiva o elimina lógicamente la cuenta. Solo protege al superadministrador y a la sesión actual; no revisa correcciones pendientes del tramitador.",
                "Riesgo de dejar trabajo asignado.",
            ),
            (
                "Eliminar la Persona",
                "Normalmente se bloquea si tiene trámites o relaciones. Cuando no hay bloqueos, conserva la ficha y la marca INACTIVA; también inactiva su Usuario.",
                "No reemplaza la baja.",
            ),
            (
                "Eliminar la Empresa",
                "Se bloquea si tiene Responsables visibles; si no, elimina lógicamente la Empresa. No es un cierre operativo completo del historial.",
                "No usar para retirar tramitadores.",
            ),
            (
                "Borrar con SQL directo",
                "Omite las reglas, transferencias, notificaciones y validaciones de la aplicación.",
                "No recomendado.",
            ),
        ],
        [2180, 4910, 2270],
        header_fill=LIGHT_AMBER,
    )
    add_callout(
        document,
        "Regla operativa",
        "Primero se da de baja la relación desde Tramitadores. Solo después, si realmente corresponde desactivar "
        "la cuenta global de la persona, se revisan todas sus relaciones, trámites y seguimientos en todas las empresas.",
        LIGHT_RED,
        RED,
    )

    add_heading(document, "6. Procedimiento seguro para una empresa", 1)
    add_heading(document, "6.1. Antes de la baja", 2)
    checks = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "Verificar que se seleccionó la empresa y el tramitador correctos.",
        "Confirmar que la cuenta principal de la empresa está activa, especialmente si existen trámites OBSERVADOS.",
        "Revisar cuántas correcciones abiertas tiene asignadas el tramitador.",
        "Si habrá reemplazo, registrar al nuevo tramitador y comprobar que tenga Usuario activo y permisos.",
        "Comprobar si la misma persona trabaja para otras empresas; su Usuario no debe desactivarse globalmente sin esa revisión.",
    ]:
        add_list_item(document, item, checks)

    add_heading(document, "6.2. Ejecutar la baja", 2)
    steps = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "Ingresar con la cuenta de la empresa.",
        "Abrir el módulo Tramitadores.",
        "Ubicar el Responsable activo que se retirará.",
        "Usar la acción Dar de baja y confirmar.",
        "Leer el mensaje final: debe indicar si no había pendientes o cuántos trámites fueron transferidos.",
    ]:
        add_list_item(document, item, steps)

    add_heading(document, "6.3. Verificaciones posteriores", 2)
    checks = add_numbering_definition(document, fmt="decimal", text="%1.", color=BLUE)
    for item in [
        "Confirmar que el Responsable figure INACTIVO y tenga fecha de baja.",
        "Entrar con la cuenta de la empresa y comprobar que las correcciones transferidas estén disponibles.",
        "Confirmar que el tramitador anterior ya no pueda iniciar ni corregir trámites para esa empresa.",
        "Comprobar que los demás tramitadores sigan activos.",
        "Abrir un trámite histórico y verificar que conserva al tramitador original.",
        "Si se desactivará el Usuario del tramitador, revisar primero otras empresas, trámites propios y seguimientos pendientes.",
    ]:
        add_list_item(document, item, checks)

    add_heading(document, "6.4. Reemplazo del tramitador", 2)
    add_body(
        document,
        "El sistema no reasigna automáticamente las correcciones pendientes al reemplazo. La baja las envía a la empresa beneficiaria. "
        "Por ello, el orden seguro es: registrar y habilitar al reemplazo, dar de baja al anterior, atender desde la cuenta de la empresa "
        "las correcciones transferidas y usar al nuevo tramitador en los trámites posteriores.",
    )

    add_heading(document, "7. Observaciones del funcionamiento actual", 1)
    add_body(
        document,
        "Estas observaciones describen riesgos encontrados en el código vigente. No significan que ya hayan sido corregidos ni que el documento haya modificado el sistema.",
    )
    add_content_table(
        document,
        ["Observación", "Impacto práctico", "Mejora sugerida"],
        [
            (
                "Registrar un tramitador no crea su User.",
                "El registro puede quedar visible, pero sin acceso operativo.",
                "Integrar la creación o vinculación de cuenta en el mismo flujo, con validación explícita.",
            ),
            (
                "Editar la ficha de empresa puede reconstruir Responsables.",
                "Un tramitador omitido puede desaparecer lógicamente sin transferir correcciones.",
                "Obligar a que toda baja use GestionTramitadoresService.",
            ),
            (
                "Inactivar o eliminar User no revisa asignaciones operativas.",
                "Puede quedar un Seguimiento activo apuntando a una cuenta inactiva.",
                "Bloquear o transferir antes de desactivar la cuenta.",
            ),
            (
                "La bandeja Mis trámites muestra trámites de la empresa a todo tramitador activo.",
                "El detalle aplica una regla más estricta: tramitador original o usuario actualmente asignado.",
                "Alinear el filtro de la bandeja con la autorización de detalle.",
            ),
            (
                "Eliminar Empresa revisa Responsables, no todo el historial del beneficiario.",
                "Puede perderse la relación Empresa visible aunque la Persona conserve trámites.",
                "Aplicar un cierre/inactivación integral y bloquear si hay historial operativo.",
            ),
        ],
        [2680, 3300, 3380],
    )

    add_heading(document, "8. Matriz de decisión rápida", 1)
    add_content_table(
        document,
        ["Necesidad", "Acción correcta"],
        [
            ("El tramitador deja de trabajar con una empresa", "Usar Tramitadores > Dar de baja."),
            ("Hay correcciones pendientes", "Asegurar primero que la empresa tenga Usuario activo; luego dar de baja."),
            ("Habrá un nuevo tramitador", "Registrar y habilitar al nuevo; después dar de baja al anterior."),
            ("La persona trabaja con otra empresa", "No desactivar su Usuario global por la baja de una sola relación."),
            ("Solo se quiere impedir el acceso total", "Primero dar de baja todas las relaciones y resolver seguimientos; después evaluar el Usuario."),
            ("Se quiere conservar auditoría", "No borrar Personas ni trámites; conservar los ids históricos del certificado."),
        ],
        [3440, 5920],
        header_fill=LIGHT_TEAL,
    )

    add_heading(document, "9. Archivos del sistema revisados", 1)
    add_body(
        document,
        "La explicación se basó en la implementación presente en los siguientes archivos:",
    )
    sources = [
        "app/Services/GestionTramitadoresService.php",
        "app/Http/Controllers/TramitadorController.php",
        "app/Http/Controllers/SeguimientoController.php",
        "app/Http/Controllers/ResponsableController.php",
        "app/Http/Controllers/PersonaController.php",
        "app/Http/Controllers/EmpresaController.php",
        "app/Http/Controllers/UsuarioController.php",
        "app/Livewire/Datatables/SeguimientoTable.php",
        "app/Models/Responsable.php",
        "app/Models/Empresa.php",
        "app/Models/Persona.php",
        "app/Models/Certificado.php",
        "app/Models/Seguimiento.php",
        "routes/admin.php",
    ]
    bullets = add_numbering_definition(document, fmt="bullet", text="•", color=TEAL)
    for source in sources:
        add_list_item(document, source, bullets, color=SLATE)

    add_callout(
        document,
        "Conclusión",
        "La baja diseñada para Tramitadores es una baja de relación, no una eliminación de identidad ni de historial. "
        "Su objetivo es retirar el acceso operativo de una empresa concreta y evitar que las correcciones abiertas queden sin responsable.",
        LIGHT_BLUE,
        BLUE,
    )

    document.core_properties.title = "Guía de flujo de empresa y tramitadores - Sistema Certificador INSO"
    document.core_properties.subject = "Funcionamiento actual, baja segura y manejo de trámites"
    document.core_properties.author = "Codex - revisión del código del sistema"
    document.core_properties.keywords = "INSO, empresa, tramitador, trámites, baja, flujo"
    document.core_properties.comments = "Documento generado a partir de la implementación vigente al 30/07/2026."

    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    result = build_document()
    print(result)
